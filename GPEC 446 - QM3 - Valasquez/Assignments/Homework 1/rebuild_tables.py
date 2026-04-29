#!/usr/bin/env python3
"""
Rebuild the 4 regression tables in HW1_Agunias.docx.

The voice-pass pipeline flattened stargazer HTML tables into stacks of
single-cell paragraphs. This script:
  1. Defines the 4 tables as Python data structures (coefficients + SEs
     read straight from the broken paragraph sequence, so the numbers
     are identical to what stargazer produced in R).
  2. Locates the broken-table paragraph ranges by anchor text (table title
     to last note line).
  3. Replaces each range with a single proper Word table.

Run:  python3 rebuild_tables.py
"""

from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT


def apply_borders(tbl):
    """Attach single-line black borders to every side + inside of the table."""
    tblPr = tbl._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._element.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)

SRC = "HW1_Agunias.docx"
OUT = "HW1_Agunias.docx"  # overwrite in place; broken copy already backed up

# -------------------------------------------------------------------
# Table data. Every coefficient/SE below is copied verbatim from the
# flattened text in the current broken docx (which preserved the
# stargazer numbers). Significance stars are preserved.
# -------------------------------------------------------------------

TABLES = [
    # -------- Table A: Q3 four-city simple regressions --------
    {
        "anchor_first": "Poverty and Mobility: Simple Regressions by City",
        "anchor_last":  "Coefficient = 1000 change in mobility per 1pp of poverty.",
        "title":        "Poverty and Mobility: Simple Regressions by City",
        "dep_var":      "Child Income at p25 ($1000s)",
        "col_labels":   ["Los Angeles", "New York", "Chicago", "New Orleans"],
        "col_numbers":  ["(1)", "(2)", "(3)", "(4)"],
        "rows": [
            # (label, [coef cells])   SE row follows in parentheses
            ("Poverty Rate 1990 (pp)", ["-0.333***", "-0.411***", "-0.455***", "-0.339***"],
                                       ["(0.009)",   "(0.011)",   "(0.012)",   "(0.017)"]),
            ("Constant",               ["40.566***", "45.219***", "40.061***", "38.267***"],
                                       ["(0.151)",   "(0.218)",   "(0.226)",   "(0.480)"]),
        ],
        "stats": [
            ("Observations", ["3,881", "2,932", "1,986", "405"]),
            ("R2",           ["0.256", "0.337", "0.418", "0.493"]),
            ("Adjusted R2",  ["0.256", "0.337", "0.418", "0.491"]),
        ],
        "notes": [
            "Note: *p<0.1; **p<0.05; ***p<0.01",
            "One simple OLS per city.",
            "Coefficient = $1000 change in mobility per 1pp of poverty.",
        ],
    },
    # -------- Table B: Q5 New Orleans short/long/interaction --------
    {
        "anchor_first": "New Orleans: Poverty, Race, and Upward Mobility",
        "anchor_last":  "Coefficient units: 1000s per pp of poverty.",
        "title":        "New Orleans: Poverty, Race, and Upward Mobility",
        "dep_var":      "Child Income at p25 ($1000s)",
        "col_labels":   ["Short (Q3)", "Long (Q4)", "Interaction (Q5)"],
        "col_numbers":  ["(1)", "(2)", "(3)"],
        "rows": [
            ("Poverty Rate 1990 (pp)",          ["-0.339***", "-0.219***", "-0.442***"],
                                                ["(0.017)",   "(0.019)",   "(0.035)"]),
            ("Majority Non-White (= 1)",        ["",          "-6.704***", "-12.727***"],
                                                ["",          "(0.656)",   "(1.007)"]),
            ("Poverty x Majority Non-White",    ["",          "",          "0.306***"],
                                                ["",          "",          "(0.041)"]),
            ("Constant",                        ["38.267***", "38.305***", "41.431***"],
                                                ["(0.480)",   "(0.428)",   "(0.577)"]),
        ],
        "stats": [
            ("Observations", ["405", "405", "405"]),
            ("R2",           ["0.493", "0.597", "0.647"]),
            ("Adjusted R2",  ["0.491", "0.595", "0.645"]),
        ],
        "notes": [
            "Note: *p<0.1; **p<0.05; ***p<0.01",
            "Sample: New Orleans commuting zone tracts.",
            "Coefficient units: $1000s per pp of poverty.",
        ],
    },
    # -------- Table C: Q7-9 Pooled OLS vs CZ FE --------
    {
        "anchor_first": "National Sample: Pooled OLS vs. Commuting-Zone FE",
        "anchor_last":  "Identification in column (2) is within-CZ only.",
        "title":        "National Sample: Pooled OLS vs. Commuting-Zone FE",
        "dep_var":      "Child Income at p25 ($1000s)",
        "col_labels":   ["Pooled OLS", "+ CZ Fixed Effects"],
        "col_numbers":  ["(1)", "(2)"],
        "rows": [
            ("Poverty Rate 1990 (pp)",   ["-0.293***", "-0.260***"], ["(0.003)", "(0.002)"]),
            ("Majority Non-White (= 1)", ["-3.310***", "-5.614***"], ["(0.070)", "(0.069)"]),
            ("Constant",                 ["39.078***", "32.047***"], ["(0.039)", "(0.486)"]),
        ],
        "stats": [
            ("Commuting-zone FE", ["No",     "Yes"]),
            ("Observations",      ["71,922", "71,921"]),
            ("R2",                ["0.274",  "0.503"]),
            ("Adjusted R2",       ["0.274",  "0.498"]),
        ],
        "notes": [
            "Note: *p<0.1; **p<0.05; ***p<0.01",
            "CZ fixed effects absorb all between-city variation.",
            "Identification in column (2) is within-CZ only.",
        ],
    },
    # -------- Table D: Q11 IV --------
    {
        "anchor_first": "Angrist-Evans IV: Same-Sex Instrument for Third Child",
        "anchor_last":  "Use ivreg() for inference.",
        "title":        "Angrist-Evans IV: Same-Sex Instrument for Third Child",
        "dep_var":      "Dep. var: morekids_num (col 1); work (cols 2-4)",
        "col_labels":   ["1st Stage (OLS)", "Reduced Form (OLS)", "Manual 2SLS (OLS)", "AER ivreg()"],
        "col_numbers":  ["(1)", "(2)", "(3)", "(4)"],
        "rows": [
            ("Same Gender (instrument)",        ["0.067***", "-0.403",  "",        ""],
                                                ["(0.006)",  "(0.253)", "",        ""]),
            ("Predicted More Kids (manual 2SLS)", ["",       "",        "-6.033",  ""],
                                                  ["",       "",        "(3.792)", ""]),
            ("More Kids (IV)",                  ["",         "",        "",        "-6.033"],
                                                ["",         "",        "",        "(3.758)"]),
            ("Constant",                        ["0.344***", "19.412***", "21.488***", "21.488***"],
                                                ["(0.004)",  "(0.180)",   "(1.437)",   "(1.425)"]),
        ],
        "stats": [
            ("Observations", ["30,000", "30,000", "30,000", "30,000"]),
            ("R2",           ["0.005",  "0.0001", "0.0001", "0.018"]),
            ("Adjusted R2",  ["0.005",  "0.0001", "0.0001", "0.018"]),
        ],
        "notes": [
            "Note: *p<0.1; **p<0.05; ***p<0.01",
            "Manual 2SLS point estimate = ivreg() point estimate.",
            "Manual 2SLS standard errors are understated.",
            "Use ivreg() for inference.",
        ],
    },
]


# -------------------------------------------------------------------
# Helpers
# -------------------------------------------------------------------

def p_text(p_element):
    return "".join(t.text or "" for t in p_element.iter(qn("w:t")))


def set_cell_text(cell, text, *, bold=False, italic=False, align=None, size=None):
    """Replace cell contents with a single paragraph of styled text."""
    # Clear existing paragraphs
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)


def build_table(doc, spec):
    """Construct a regression table and return its OOXML element (unattached)."""
    ncols = 1 + len(spec["col_labels"])
    # rows: title / dep_var / col_labels / col_numbers / coef+SE pairs / stats / notes
    nrows = 4 + 2 * len(spec["rows"]) + len(spec["stats"]) + len(spec["notes"])

    tbl = doc.add_table(rows=nrows, cols=ncols)
    apply_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    LEFT   = WD_ALIGN_PARAGRAPH.LEFT

    r = 0
    # --- Row 0: title (merged) ---
    title_row = tbl.rows[r].cells
    merged = title_row[0]
    for c in title_row[1:]:
        merged = merged.merge(c)
    set_cell_text(merged, spec["title"], bold=True, align=CENTER, size=11)
    r += 1

    # --- Row 1: "Dependent variable: X" merged over data cols ---
    dep_row = tbl.rows[r].cells
    set_cell_text(dep_row[0], "")
    merged = dep_row[1]
    for c in dep_row[2:]:
        merged = merged.merge(c)
    set_cell_text(merged, f"Dependent variable: {spec['dep_var']}", italic=True, align=CENTER, size=10)
    r += 1

    # --- Row 2: column labels ---
    lab_row = tbl.rows[r].cells
    set_cell_text(lab_row[0], "")
    for i, lbl in enumerate(spec["col_labels"], start=1):
        set_cell_text(lab_row[i], lbl, bold=True, align=CENTER, size=10)
    r += 1

    # --- Row 3: column numbers (1) (2) ... ---
    num_row = tbl.rows[r].cells
    set_cell_text(num_row[0], "")
    for i, n in enumerate(spec["col_numbers"], start=1):
        set_cell_text(num_row[i], n, align=CENTER, size=10)
    r += 1

    # --- coef + SE rows ---
    for label, coefs, ses in spec["rows"]:
        coef_row = tbl.rows[r].cells
        set_cell_text(coef_row[0], label, align=LEFT, size=10)
        for i, v in enumerate(coefs, start=1):
            set_cell_text(coef_row[i], v, align=CENTER, size=10)
        r += 1
        se_row = tbl.rows[r].cells
        set_cell_text(se_row[0], "", size=10)
        for i, v in enumerate(ses, start=1):
            set_cell_text(se_row[i], v, align=CENTER, size=10)
        r += 1

    # --- stats rows ---
    for label, vals in spec["stats"]:
        row = tbl.rows[r].cells
        set_cell_text(row[0], label, align=LEFT, size=10)
        for i, v in enumerate(vals, start=1):
            set_cell_text(row[i], v, align=CENTER, size=10)
        r += 1

    # --- notes rows (each merged across all cols) ---
    for note in spec["notes"]:
        nrow = tbl.rows[r].cells
        merged = nrow[0]
        for c in nrow[1:]:
            merged = merged.merge(c)
        set_cell_text(merged, note, italic=True, align=LEFT, size=9)
        r += 1

    # Detach from document's append location so caller can place it precisely
    element = tbl._element
    element.getparent().remove(element)
    return element


def find_para_range(body_children, anchor_first, anchor_last):
    """Return (start_idx, end_idx_inclusive) into body_children for paragraphs
    whose visible text exactly matches anchor_first then later anchor_last."""
    start = None
    end = None
    for i, ch in enumerate(body_children):
        if ch.tag != qn("w:p"):
            continue
        txt = p_text(ch).strip().replace("\xa0", " ")
        if start is None and txt == anchor_first:
            start = i
        elif start is not None and txt == anchor_last:
            end = i
            break
    if start is None or end is None:
        raise RuntimeError(f"Could not locate range [{anchor_first!r} .. {anchor_last!r}] "
                           f"(start={start}, end={end})")
    return start, end


def main():
    doc = Document(SRC)
    body = doc.element.body

    for spec in TABLES:
        # Fresh snapshot of children each iteration because we mutate the tree.
        children = list(body.iterchildren())
        start, end = find_para_range(children, spec["anchor_first"], spec["anchor_last"])
        print(f"Replacing [{spec['anchor_first'][:50]}...]: child indices {start}..{end} "
              f"({end - start + 1} paragraphs)")

        # Build the table element (detached).
        tbl_element = build_table(doc, spec)

        # Insert the table immediately before the first broken paragraph.
        first_broken = children[start]
        first_broken.addprevious(tbl_element)

        # Also insert an empty paragraph after the table for spacing, so the
        # following heading does not collide with the table.
        spacer_p = doc.paragraphs[-1]._element  # grabs something; safer to build fresh
        # Build a fresh empty paragraph element via python-docx then detach.
        tmp = doc.add_paragraph("")
        spacer = tmp._element
        spacer.getparent().remove(spacer)
        first_broken.addprevious(spacer)

        # Remove the broken paragraphs (start..end inclusive).
        for ch in children[start:end + 1]:
            ch.getparent().remove(ch)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
