"""Build the ORANGE Memo (Myanmar) as a properly formatted .docx.

Format per syllabus (pp. 1-2): 3 pages max, double-spaced, 12pt, 1-inch margins,
footnote citations only (no parentheticals). Target 950 words body.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def add_footnote(paragraph, footnote_text, footnote_id, doc):
    """Append a footnote reference to the end of `paragraph` and register the
    footnote body text. Uses the document's footnotes part, creating it if
    absent."""
    run = paragraph.add_run()
    # Footnote reference character run with footnoteReference element
    ref = OxmlElement('w:footnoteReference')
    ref.set(qn('w:id'), str(footnote_id))
    rpr = OxmlElement('w:rPr')
    rstyle = OxmlElement('w:rStyle')
    rstyle.set(qn('w:val'), 'FootnoteReference')
    rpr.append(rstyle)
    run._r.append(rpr)
    run._r.append(ref)


def main():
    doc = Document()

    # Page margins — 1 inch all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default style: Times New Roman, 12pt, double-spaced
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Header block (single-spaced, tight) — memo-style heading
    header_lines = [
        ("Edgar Agunias", True),
        ("GPCO 410 — International Politics and Security", False),
        ("Professor Prather", False),
        ("Orange Memo — Myanmar, 29 April 2026", False),
    ]
    for text, bold in header_lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = bold

    # Title
    title = doc.add_paragraph()
    title.alignment = 1  # center
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    tr = title.add_run("Democratization without Disarmament: Why Myanmar’s 2010–2021 Arc Raised Civil War Risk")
    tr.bold = True

    # Body paragraphs. Footnote numerals appear as superscripted digits in
    # brackets for drafting clarity; Edgar will convert to Word footnotes in
    # the voice pass (python-docx does not ship first-class footnote support
    # without manipulating the footnotes part directly, which is fragile).
    # Superscript numerals are rendered via character runs.

    def add_body(text_runs):
        """text_runs is a list of (text, is_superscript) tuples."""
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        for text, sup in text_runs:
            r = p.add_run(text)
            if sup:
                r.font.superscript = True
        return p

    # Paragraph 1 — Answer + frame (~90w)
    add_body([
        ("Myanmar was at greater risk of political violence and civil war across the full 2010–2021 arc, and the February 2021 coup is the predicted realization of that risk rather than a departure from it. The 2010 opening produced the kind of incomplete democratization Cederman, Hug, and Krebs identify as a governmental-conflict risk zone,", False),
        ("1", True),
        (" while the 2008 constitution preserved a Walter-style commitment problem the Tatmadaw refused to surrender.", False),
        ("2", True),
        (" When the 2020 election made that reservation untenable, the military struck, and the resulting war is the one the theory expects.", False),
    ])

    # Paragraph 2 — Theory (~140w)
    add_body([
        ("Two claims from the theoretical literature carry the argument. First, Cederman, Hug, and Krebs find that periods of democratization — not Polity levels — significantly raise the probability of civil war onset, and that the effect is concentrated in governmental, center-seeking conflicts rather than territorial ones.", False),
        ("3", True),
        (" Their mechanism is elite competition: transitions mobilize mass audiences around state capture faster than institutions can absorb the contestation. Second, Walter argues that bargaining failures produce civil war most often through commitment problems, and that regime transitions are the paradigmatic case because no third party can credibly bind either side to honor a future distribution of power.", False),
        ("4", True),
        (" Powell provides the underlying vocabulary: shifting power plus non-binding promises make preventive action rational for the declining side.", False),
        ("5", True),
        (" Myanmar 2010–2021 runs on both rails at once — a period of change and an unresolved commitment problem embedded in the founding constitutional bargain.", False),
    ])

    # Paragraph 3 — Incomplete democratization, Tatmadaw strategic choice (~200w)
    add_body([
        ("The 2010 opening was incomplete democratization by design. The Tatmadaw, facing the choice among continued direct military rule, genuine civilian transfer, or a managed hybrid, chose the third. Continued junta rule was dominated by rising domestic and sanctions costs after 2007, and genuine transfer was dominated by the loss of prosecution immunity and economic holdings. The chosen strategy — the 2008 constitution’s “disciplined democracy” — reserved 25 percent of parliamentary seats for serving officers under Articles 109(b) and 141(b), placed the Home, Defense, and Border ministries under the Commander-in-Chief, and required a greater-than-75-percent threshold for constitutional amendment under Article 436.", False),
        ("6", True),
        (" This is a strategy set chosen deliberately: the Tatmadaw preferred a veto architecture to either pole because it preserved autonomy, coercive control, and a future coup option at the lowest expected cost. Its belief — that reserved seats plus ministerial control would hold indefinitely against any civilian majority — made the gamble look safe. Yet under Cederman this institutional configuration sits squarely in the governmental-conflict risk band, and under Walter it is a commitment problem parked in constitutional text rather than resolved by it. The 2015 NLD landslide exposed the gap: an elected civilian government now presided over a state whose coercive instruments it did not control.", False),
    ])

    # Paragraph 4 — Rohingya proof (~70–90w)
    add_body([
        ("The 2016–2017 Rakhine clearance operations confirmed that the reform period had not bound the Tatmadaw. The military executed operations displacing roughly 740,000 Rohingya to Bangladesh without civilian authorization and over civilian reputational cost, and Aung San Suu Kyi’s 2019 defense of the Tatmadaw at the International Court of Justice demonstrated how dependent her government was on the generals it could not discipline.", False),
        ("7", True),
        (" If reforms had actually domesticated military coercion, this pattern would have been harder to sustain.", False),
    ])

    # Paragraph 5 — 2020 trigger, beliefs, payoffs, no-settlement beat (~220w)
    add_body([
        ("The 2020 election made the constitutional bargain’s future cost intolerable to the Tatmadaw and is the memo’s analytical crux. The NLD won 82 percent of contested lower-house seats, up from 77 percent in 2015, and began signaling a sustained push for amendment of Articles 436 and the reserved-seats provisions.", False),
        ("8", True),
        (" The Tatmadaw faced three strategies: accept the result and ride out a shrinking future share, renegotiate the constitutional settlement while leverage remained, or preempt through coup under Article 417’s emergency clause.", False),
        ("9", True),
        (" Acceptance was rejected because the Tatmadaw believed the 2020 margin signaled permanent electoral lockout and that the NLD, once able to assemble a cross-bench supermajority, would not credibly refrain from amendment, prosecution, or economic reprisal. Renegotiation was rejected because no NLD commitment could bind future civilian majorities — the textbook Walter commitment problem. The NLD, believing international legitimacy would deter a coup, chose to push rather than compromise. The expected payoff of coup-now dominated: immediate coercive supremacy at known cost, versus accepting a distribution of power trending toward the Tatmadaw’s juridical exposure. A settlement — amendment of Article 436 in exchange for guaranteed military carve-outs and prosecution immunity — was theoretically available, but went unsigned because neither side could credibly enforce the other’s forbearance across future elections. Fraud allegations supplied the pretext; the commitment problem supplied the cause.", False),
    ])

    # Paragraph 6 — Post-coup + rebuttal + implication (~200w)
    add_body([
        ("Post-coup Myanmar is the predicted realization. The NUG declared in April 2021 and the PDFs mobilized nationally from May, pulling fighting into Bamar-majority heartland regions for the first time in a generation; the Three Brotherhood Alliance’s Operation 1027 in October 2023 then expanded the war into a coordinated, multi-front campaign that reduced junta township control below half the country.", False),
        ("10", True),
        (" ACLED records more than 50,000 reported fatalities and UNOCHA over three million internally displaced by late 2024.", False),
        ("11", True),
        (" Critically, the post-coup war is governmental in character — the PDFs seek to restore civilian rule at the center — which is precisely the Cederman prediction for democratization-era conflict. The strongest objection is the stability narrative: 2010–2020 looked calmer than the decades before it, ceasefires multiplied, and the economy opened. But that calm rested on the Tatmadaw’s intact veto, not a resolved bargain; the 2021 coup revealed the equilibrium as a frozen pause in an unenforced commitment problem, not durable peace. The generalizable lesson is that elections without civilian control of coercion build commitment problems rather than democracies, and raise rather than lower the probability of war.", False),
    ])

    # Closing sentence (~25w) — kept as its own paragraph
    add_body([
        ("Myanmar did not grow safer by democratizing and then violent after the coup; it grew violent because its democratization never removed the military’s coercive veto in the first place.", False),
    ])

    # ------- FOOTNOTES section (appended as endnotes-style list for Edgar to
    # convert into real Word footnotes during the voice pass). This is the
    # pragmatic route given python-docx's limited footnote support.
    fn_head = doc.add_paragraph()
    fn_head.paragraph_format.space_before = Pt(18)
    fn_head.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = fn_head.add_run("Footnotes (convert to Word footnotes during voice pass)")
    r.bold = True
    r.italic = True

    footnotes = [
        "Lars-Erik Cederman, Simon Hug, and Lutz F. Krebs, “Democratization and Civil War: Empirical Evidence,” Journal of Peace Research 47, no. 4 (2010): 377–94.",
        "Barbara F. Walter, “Bargaining Failures and Civil War,” Annual Review of Political Science 12 (2009): 243–61.",
        "Cederman, Hug, and Krebs, “Democratization and Civil War,” 383, Table I; 387 on the governmental-versus-territorial distinction.",
        "Walter, “Bargaining Failures and Civil War,” 248–52 on commitment problems in regime transitions.",
        "Robert Powell, “War as a Commitment Problem,” International Organization 60, no. 1 (2006): 169–203.",
        "Constitution of the Republic of the Union of Myanmar (2008), arts. 109(b), 141(b), 436.",
        "Office of the United Nations High Commissioner for Refugees, “Rohingya Emergency,” operational update, 2018; Independent International Fact-Finding Mission on Myanmar, Report of the Detailed Findings, A/HRC/39/CRP.2 (17 September 2018); International Court of Justice, The Gambia v. Myanmar, verbatim record CR 2019/18 (11 December 2019).",
        "Union Election Commission of Myanmar, 2020 general election results; International Crisis Group, Responding to the Myanmar Coup, Asia Briefing No. 166 (16 February 2021).",
        "Constitution of the Republic of the Union of Myanmar (2008), art. 417.",
        "Special Advisory Council for Myanmar, Effective Control in Myanmar (September 2023); International Crisis Group, A Silent Revolution: Myanmar’s Junta in Crisis, Asia Report No. 334 (November 2024).",
        "ACLED, Myanmar Conflict Dashboard (accessed 2026); United Nations Office for the Coordination of Humanitarian Affairs, Myanmar Humanitarian Needs and Response Plan 2024.",
    ]
    for i, fn in enumerate(footnotes, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(f"{i}. ")
        r.bold = True
        p.add_run(fn)

    # ------- Disclosure block
    doc.add_paragraph()
    disc_head = doc.add_paragraph()
    disc_head.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = disc_head.add_run("— Output Disclosure —")
    r.bold = True
    for line in [
        "Generated for: Edgar Agunias",
        "Date: 2026-04-24",
        "Model: Claude Opus 4.7 (1M context)",
        "Agent: Athena (GPCO 410)",
        "Sources: orange_memo_kickoff.md, outline.md, reading_notes.md, myanmar_context.md, prompt.md; Claudia/COURSE_MEMORY.md; Claudia/FEEDBACK.md (2026-04-14, 2026-04-23).",
        "Status: First full draft. Requires Edgar voice pass, footnote conversion to live Word footnotes, word-count trim to 950-word graded ceiling, and the ai-disclosure block per _claudia/sop/ai-disclosure.md to be pasted in before Canvas upload.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.add_run(line)

    out_path = "/Users/edgar/Documents/01 Projects/Claudia/GPCO 410 - Intl Pol:Sec - Praether/Assignments/Orange Memo - Myanmar/Orange_Memo_Myanmar_DRAFT.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
